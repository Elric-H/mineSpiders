import os
import traceback
from docx import Document
from docxcompose.composer import Composer
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

class SummaryGenerator:
    def __init__(self, output_path="合并结果.docx"):
        self.output_path = os.path.abspath(output_path)

    def _insert_section_break(self, doc):
        """在文档末尾插入分节符 (section break)"""
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'section')  # 设置为分节符
        run._r.append(br)

    def generate(self, source_files):
        if not source_files:
            print("错误：未提供任何文件")
            return False

        valid_files = [os.path.abspath(f) for f in source_files if os.path.exists(f)]
        if not valid_files:
            print("错误：没有有效的可合并文件")
            return False

        try:
            print(f"共检测到 {len(valid_files)} 个有效文件。")

            # 以第一个文档作为主文档
            master_doc = Document(valid_files[0])
            composer = Composer(master_doc)

            for idx, file in enumerate(valid_files):
                if idx == 0:
                    continue  # 第一个文档已加载为 master
                print(f"正在合并 ({idx + 1}/{len(valid_files)}): {os.path.basename(file)}")
                try:
                    # 合并前在主文档末尾插入分节符
                    self._insert_section_break(composer.doc)

                    sub_doc = Document(file)
                    composer.append(sub_doc)

                except Exception as e:
                    print(f"⚠️ 读取失败，跳过: {file} - {e}")
                    continue

            composer.save(self.output_path)
            print(f"✅ 合并完成！输出文件：{self.output_path}")
            return True

        except Exception as e:
            print(f"❌ 合并失败: {str(e)}")
            traceback.print_exc()
            return False
