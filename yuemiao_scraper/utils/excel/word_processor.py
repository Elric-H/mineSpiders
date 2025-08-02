from docx import Document
from typing import Dict, Any
import traceback
from datetime import datetime

from summary_generator import SummaryGenerator


class WordProcessor:
    def __init__(self, template_path: str,
                 summary_enabled: bool = True,
                 summary_filename: str = "./汇总.docx",
                 auto_generate_summary: bool = True):
        self.template_path = template_path
        self.summary_enabled = summary_enabled
        self.summary_generator = SummaryGenerator(summary_filename)  # 使用独立的汇总生成器
        self.auto_generate_summary = auto_generate_summary
        self.generated_files = []

    def generate_document(self, output_path: str, data: Dict[str, Any]) -> bool:
        try:
            doc = Document(self.template_path)
            self._replace_placeholders(doc, data)
            self._process_row_data(doc, data.get('rows', []))
            doc.save(output_path)

            if self.summary_enabled:
                self.generated_files.append(output_path)
                if self.auto_generate_summary:
                    self.summary_generator.generate(self.generated_files)

            return True
        except Exception as e:
            print(f"生成文档时出错: {str(e)}")
            traceback.print_exc()
            return False

    def generate_summary(self) -> bool:
        if not self.summary_enabled or not self.generated_files:
            return False

        try:
            summary_doc = Document()

            # 添加汇总文档标题
            summary_doc.add_heading('文档内容汇总', level=0)
            summary_doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            summary_doc.add_paragraph(f"共合并 {len(self.generated_files)} 个文档:")
            summary_doc.add_page_break()

            # 合并所有文档内容
            for file_path in self.generated_files:
                try:
                    # 添加文档来源标识
                    summary_doc.add_heading(f"文档来源: {file_path}", level=1)

                    # 读取源文档内容
                    src_doc = Document(file_path)

                    # 复制所有段落
                    for para in src_doc.paragraphs:
                        new_para = summary_doc.add_paragraph()
                        for run in para.runs:
                            new_run = new_para.add_run(run.text)
                            new_run.font.name = run.font.name
                            new_run.font.size = run.font.size
                            new_run.font.bold = run.font.bold
                            new_run.font.italic = run.font.italic
                            new_run.font.underline = run.font.underline
                            if run.font.color.rgb:
                                new_run.font.color.rgb = run.font.color.rgb

                    # 复制所有表格
                    for table in src_doc.tables:
                        new_table = summary_doc.add_table(rows=1, cols=len(table.columns))
                        # 复制表头
                        for i, cell in enumerate(table.row_cells(0)):
                            new_table.cell(0, i).text = cell.text
                        # 复制数据行
                        for row in table.rows[1:]:
                            new_row = new_table.add_row()
                            for i, cell in enumerate(row.cells):
                                new_row.cells[i].text = cell.text

                    summary_doc.add_page_break()
                except Exception as e:
                    print(f"合并文档 {file_path} 时出错: {str(e)}")
                    continue

            summary_doc.save(self.summary_filename)
            print(f"内容汇总文档已生成: {self.summary_filename}")
            return True
        except Exception as e:
            print(f"生成内容汇总文档时出错: {str(e)}")
            traceback.print_exc()
            return False

    def _replace_placeholders(self, doc, data):
        replacements = {
            '${name}': data.get('a', ''),
            '${code}': data.get('b', ''),
            '${dept}': data.get('c', ''),
            '${post}': data.get('d', ''),
            '${date}': datetime.now().strftime('%Y-%m-%d')
        }

        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_text, new_text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in replacements.items():
                            if old_text in paragraph.text:
                                paragraph.text = paragraph.text.replace(old_text, new_text)

    def _process_row_data(self, doc, rows_data):
        if not rows_data:
            return

        for table_idx, table in enumerate(doc.tables, 1):
            for row_idx, row in enumerate(table.rows):
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if '${row}' in paragraph.text:
                            # 1. 保存原行的样式
                            template_row = row

                            # 2. 删除占位符
                            paragraph.text = paragraph.text.replace('${row}', '')

                            # 3. 插入带样式的新行
                            self._insert_rows_after(table, row_idx + 1, template_row, rows_data)

                            # 4. 删除原模板行
                            table._tbl.remove(row._tr)
                            return

    def _insert_rows_after(self, table, target_row_idx, template_row, rows_data):
        """在指定行后插入带样式和数据的新行"""
        try:
            for i, row_data in enumerate(rows_data, 1):
                # 复制模板行的样式
                new_row = self._copy_row(table, template_row)

                # 移动新行到正确位置
                table._tbl.insert(target_row_idx + i, new_row._tr)

                # 填充数据
                for col_idx, cell in enumerate(new_row.cells):
                    if col_idx == 0:
                        cell.text = str(i)  # 序号列
                    elif col_idx - 1 < len(row_data):
                        cell.text = str(row_data[col_idx - 1]) if row_data[col_idx - 1] is not None else ''

        except Exception as e:
            print(f"插入带样式行时出错: {str(e)}")
            traceback.print_exc()

    def _copy_row(self, table, template_row):
        """复制行并设置边框样式，首尾单元格边框加粗"""
        new_row = table.add_row()

        # 复制行高
        if template_row.height is not None:
            new_row.height = template_row.height

        # 获取表格总列数
        col_count = len(new_row.cells)

        for i, cell in enumerate(new_row.cells):
            if i < len(template_row.cells):
                # 复制单元格宽度
                cell.width = template_row.cells[i].width

                # 设置边框样式
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()

                from docx.oxml.shared import OxmlElement
                from docx.oxml.ns import qn

                borders = OxmlElement('w:tcBorders')

                # 定义边框属性函数
                def add_border(border_name, size=4):
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), str(size))
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')
                    borders.append(border)

                # 普通单元格边框设置(1磅)
                # 左侧边框 - 如果是第一个单元格则加粗(2磅)
                add_border('left', 8 if i == 0 else 4)
                # 右侧边框 - 如果是最后一个单元格则加粗(2磅)
                add_border('right', 8 if i == col_count - 1 else 4)
                # 上下边框保持1磅
                add_border('top')
                add_border('bottom')

                tcPr.append(borders)

                # 设置单元格垂直居中
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')  # 垂直居中
                tcPr.append(vAlign)

                # 复制段落样式
                for paragraph in cell.paragraphs:
                    paragraph.style = template_row.cells[i].paragraphs[0].style

                    # 复制字体样式
                    if template_row.cells[i].paragraphs[0].runs:
                        run = paragraph.add_run()
                        font = template_row.cells[i].paragraphs[0].runs[0].font
                        run.font.name = font.name
                        run.font.size = font.size
                        run.font.bold = font.bold
                        run.font.italic = font.italic
                        run.font.underline = font.underline
                        run.font.color.rgb = font.color.rgb

        return new_row
